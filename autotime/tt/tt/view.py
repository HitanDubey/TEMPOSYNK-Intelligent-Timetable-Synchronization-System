## MODIFIED: Individual Faculty Timetable Generation with NAME MERGING
## Now combines all courses for same faculty name into ONE timetable

from django.shortcuts import render, redirect
from django.http import HttpResponse, FileResponse
from . import pool
import random
import io
import pandas as pd
from docx import Document
from xhtml2pdf import pisa
from collections import defaultdict
import json
from datetime import datetime
import time
import traceback

# Cache for course names to avoid repeated database queries
_course_name_cache = {}

def get_course_name(course_code):
    """Get course name from course code with caching"""
    if course_code in _course_name_cache:
        return _course_name_cache[course_code]
    
    dbe, cmd = pool.ConnectionPool()
    try:
        cmd.execute("SELECT course_name FROM Courses WHERE course_code = ?", [course_code])
        result = cmd.fetchone()
        if result:
            _course_name_cache[course_code] = result[0]
            return result[0]
        return course_code
    except Exception:
        return course_code
    finally:
        dbe.close()

def callindex(request):
    return render(request, 'index.html')

def contact(request):
    return render(request, 'contact.html')

def about(request):
    return render(request, 'about.html')

def login(request):
    return render(request, 'login.html')

def signin(request):
    return render(request, 'signin.html')

def timetable(request):
    return render(request, 'timetable.html')

def program(request):
    return render(request, 'program.html')

def faculty(request):
    return render(request, 'faculty.html')

def subject(request):
    return render(request, 'subject.html')

def lecture(request):
    return render(request, 'lecture.html')

def lab(request):
    return render(request, 'lab.html')

def generate(request):
    """Show generate page with program/semester options"""
    dbe, cmd = pool.ConnectionPool()
    try:
        cmd.execute("SELECT DISTINCT Program_name FROM Program ORDER BY Program_name")
        programs = [row[0] for row in cmd.fetchall()]
        cmd.execute("SELECT DISTINCT semester_id FROM Program ORDER BY CAST(semester_id AS INTEGER)")
        semesters = [row[0] for row in cmd.fetchall()]
    except Exception as e:
        programs = []
        semesters = []
        print(f"Error fetching data: {e}")
    finally:
        dbe.close()
    return render(request, 'generate.html', {'programs': programs, 'semesters': semesters})

def output(request):
    """Display generated timetable with individual faculty timetables"""
    schedules = request.session.get('schedule_results')
    faculty_individual_timetables = request.session.get('faculty_individual_timetables', {})
    return render(request, 'output.html', {
        'schedules': schedules,
        'faculty_individual_timetables': faculty_individual_timetables
    })

def view_course(request):
    return render(request, 'view_course.html')

def view_faculty(request):
    return render(request, 'view_faculty.html')

def view_labRoom(request):
    return render(request, 'view_labRoom.html')

def view_lectureRoom(request):
    return render(request, 'view_lectureRoom.html')

def view_program(request):
    return render(request, 'view_program.html')

def loginadmin(request):
    """Handle admin login"""
    username = request.POST.get("username", "").strip()
    password = request.POST.get("password", "").strip()

    if not username or not password:
        return render(request, 'login.html', {"msg": "Please enter both username and password!"})

    dbe, cmd = pool.ConnectionPool()
    try:
        cmd.execute("SELECT * FROM Admin WHERE username = ? AND password = ?", [username, password])
        result = cmd.fetchone()
    except Exception as e:
        result = None
        print(f"Login error: {e}")
    finally:
        dbe.close()

    if result:
        return render(request, 'timetable.html', {"msg": ""})
    else:
        return render(request, 'login.html', {"msg": "Invalid Username or Password!"})

def callsignin(request):
    """Handle user registration"""
    name = request.POST.get('name', '').strip()
    email = request.POST.get('email', '').strip()
    username = request.POST.get("username", "").strip()
    password = request.POST.get("password", "").strip()

    if not all([name, email, username, password]):
        return render(request, 'signin.html', {"msg": "Please fill all fields!"})

    if '@' not in email or '.' not in email:
        return render(request, 'signin.html', {"msg": "Please enter a valid email address!"})

    dbe, cmd = pool.ConnectionPool()
    try:
        cmd.execute("SELECT COUNT(*) FROM Admin WHERE username = ?", [username])
        count = cmd.fetchone()[0]
        
        if count > 0:
            dbe.close()
            return render(request, 'signin.html', {"msg": f"⚠️ Username '{username}' already exists! Please choose a different username."})
        
        cmd.execute(
            "INSERT INTO Admin (name, email, username, password) VALUES (?, ?, ?, ?)",
            [name, email, username, password]
        )
        dbe.commit()
        msg = "✅ Registration Successful! Please login."
    except Exception as e:
        dbe.rollback()
        msg = f"❌ Registration failed: {str(e)}"
    finally:
        dbe.close()

    return render(request, 'login.html', {"msg": msg})

def callprogram(request):
    """Add new program with duplicate check"""
    programs = request.POST.get('programs', '').strip()
    semesters = request.POST.get('semesters', '').strip()

    if not programs or not semesters:
        return render(request, 'program.html', {"msg": "⚠️ Please fill all fields!"})

    try:
        s = int(semesters)
        if s < 1 or s > 12:
            return render(request, 'program.html', {"msg": "⚠️ Semester must be between 1 and 12!"})
    except (TypeError, ValueError):
        return render(request, 'program.html', {"msg": "⚠️ Semester must be a valid number!"})

    dbe, cmd = pool.ConnectionPool()
    
    try:
        cmd.execute(
            "SELECT COUNT(*) FROM Program WHERE Program_name = ? AND semester_id = ?",
            [programs, str(s)]
        )
        count = cmd.fetchone()[0]
        
        if count > 0:
            dbe.close()
            return render(request, 'program.html', {"msg": f"⚠️ Program '{programs}' already exists for Semester {s}! Please enter a different program or semester."})
        
        cmd.execute(
            "INSERT INTO Program (Program_name, semester_id) VALUES (?, ?)",
            [programs, str(s)]
        )
        dbe.commit()
        msg = f"✅ Program '{programs}' for Semester {s} added successfully!"
            
    except Exception as e:
        dbe.rollback()
        error_msg = str(e)
        if "UNIQUE constraint failed" in error_msg:
            msg = f"⚠️ Program '{programs}' already exists for Semester {s}! Please enter a different combination."
        else:
            msg = f"❌ Error adding program: {error_msg}"
    finally:
        dbe.close()

    return render(request, 'program.html', {"msg": msg})

def callsubjects(request):
    """Add new subject/course with duplicate check"""
    course = request.POST.get('course', '').strip()
    course_n = request.POST.get('course_n', '').strip()
    frequency = request.POST.get("frequency", '').strip()
    program_name = request.POST.get("program_name", '').strip()
    semester_id = request.POST.get("semester_id", '').strip()

    if not all([course, course_n, frequency, program_name, semester_id]):
        return render(request, 'subject.html', {"msg": "⚠️ Please fill all fields!"})

    try:
        s = int(semester_id)
        freq = int(frequency)
        if freq < 1 or freq > 10:
            return render(request, 'subject.html', {"msg": "⚠️ Frequency must be between 1 and 10 hours per week!"})
        if s < 1 or s > 12:
            return render(request, 'subject.html', {"msg": "⚠️ Semester must be between 1 and 12!"})
    except (TypeError, ValueError):
        return render(request, 'subject.html', {"msg": "⚠️ Semester and Frequency must be valid numbers!"})

    dbe, cmd = pool.ConnectionPool()
    
    try:
        cmd.execute(
            "SELECT COUNT(*) FROM Program WHERE Program_name = ? AND semester_id = ?",
            [program_name, str(s)]
        )
        prog_exists = cmd.fetchone()[0]
        
        if prog_exists == 0:
            dbe.close()
            return render(request, 'subject.html', {"msg": f"⚠️ Program '{program_name}' for Semester {s} does not exist! Please add the program first."})
        
        cmd.execute(
            "SELECT COUNT(*) FROM Courses WHERE course_code = ?",
            [course]
        )
        count = cmd.fetchone()[0]
        
        if count > 0:
            dbe.close()
            return render(request, 'subject.html', {"msg": f"⚠️ Course code '{course}' already exists! Please use a different course code."})
        
        cmd.execute(
            "INSERT INTO Courses (course_code, course_name, frequency, Program_name, semester_id) VALUES (?, ?, ?, ?, ?)",
            [course, course_n, freq, program_name, s]
        )
        dbe.commit()
        msg = f"✅ Course '{course_n}' added successfully!"
            
    except Exception as e:
        dbe.rollback()
        error_msg = str(e)
        if "UNIQUE constraint failed" in error_msg:
            msg = f"⚠️ Course code '{course}' already exists! Please use a different course code."
        elif "FOREIGN KEY constraint failed" in error_msg:
            msg = f"⚠️ Program '{program_name}' for Semester {s} does not exist! Please add the program first."
        else:
            msg = f"❌ Error adding course: {error_msg}"
    finally:
        dbe.close()

    return render(request, 'subject.html', {"msg": msg})

def calllecture(request):
    """Add lecture room with duplicate check"""
    lecture_id = request.POST.get('lecture_id', '').strip()
    lecture_capacity = request.POST.get('lecture_capacity', '').strip()

    if not lecture_id or not lecture_capacity:
        return render(request, 'lecture.html', {"msg": "⚠️ Please fill all fields!"})

    try:
        s = int(lecture_capacity)
        if s < 10:
            return render(request, 'lecture.html', {"msg": "⚠️ Room capacity must be at least 10!"})
    except (TypeError, ValueError):
        return render(request, 'lecture.html', {"msg": "⚠️ Capacity must be a valid number!"})

    dbe, cmd = pool.ConnectionPool()
    
    try:
        cmd.execute(
            "SELECT COUNT(*) FROM Lecture WHERE lecture_id = ?",
            [lecture_id]
        )
        count = cmd.fetchone()[0]
        
        if count > 0:
            dbe.close()
            return render(request, 'lecture.html', {"msg": f"⚠️ Lecture room '{lecture_id}' already exists! Please use a different ID."})
        
        cmd.execute(
            "INSERT INTO Lecture (lecture_id, lecture_capacity) VALUES (?, ?)",
            [lecture_id, s]
        )
        dbe.commit()
        msg = f"✅ Lecture Room '{lecture_id}' added successfully!"
            
    except Exception as e:
        dbe.rollback()
        error_msg = str(e)
        if "UNIQUE constraint failed" in error_msg:
            msg = f"⚠️ Lecture room '{lecture_id}' already exists! Please use a different ID."
        else:
            msg = f"❌ Error adding lecture room: {error_msg}"
    finally:
        dbe.close()

    return render(request, 'lecture.html', {"msg": msg})

def calllab(request):
    """Add lab room with duplicate check"""
    lab_id = request.POST.get('lab_id', '').strip()
    lab_capacity = request.POST.get('lab_capacity', '').strip()

    if not lab_id or not lab_capacity:
        return render(request, 'lab.html', {"msg": "⚠️ Please fill all fields!"})

    try:
        s = int(lab_capacity)
        if s < 5:
            return render(request, 'lab.html', {"msg": "⚠️ Lab capacity must be at least 5!"})
    except (TypeError, ValueError):
        return render(request, 'lab.html', {"msg": "⚠️ Capacity must be a valid number!"})

    dbe, cmd = pool.ConnectionPool()
    
    try:
        cmd.execute(
            "SELECT COUNT(*) FROM Lab WHERE Lab_id = ?",
            [lab_id]
        )
        count = cmd.fetchone()[0]
        
        if count > 0:
            dbe.close()
            return render(request, 'lab.html', {"msg": f"⚠️ Lab room '{lab_id}' already exists! Please use a different ID."})
        
        cmd.execute(
            "INSERT INTO Lab (Lab_id, Lab_capacity) VALUES (?, ?)",
            [lab_id, s]
        )
        dbe.commit()
        msg = f"✅ Lab Room '{lab_id}' added successfully!"
            
    except Exception as e:
        dbe.rollback()
        error_msg = str(e)
        if "UNIQUE constraint failed" in error_msg:
            msg = f"⚠️ Lab room '{lab_id}' already exists! Please use a different ID."
        else:
            msg = f"❌ Error adding lab room: {error_msg}"
    finally:
        dbe.close()

    return render(request, 'lab.html', {"msg": msg})

def callfaculty(request):
    """Add faculty with duplicate check"""
    Professor_id = request.POST.get('Professor_id', '').strip()
    Professor_name = request.POST.get('Professor_name', '').strip()
    course_code = request.POST.get("course_code", '').strip()
    Phone_number = request.POST.get("Phone_number", '').strip()
    Email_ID = request.POST.get("Email_ID", '').strip()

    if not all([Professor_id, Professor_name, course_code, Phone_number, Email_ID]):
        return render(request, 'faculty.html', {"msg": "⚠️ Please fill all fields!"})

    try:
        s = int(Phone_number)
        if len(str(s)) != 10:
            return render(request, 'faculty.html', {"msg": "⚠️ Phone number must be exactly 10 digits!"})
    except (TypeError, ValueError):
        return render(request, 'faculty.html', {"msg": "⚠️ Phone number must be a valid number!"})

    if '@' not in Email_ID or '.' not in Email_ID:
        return render(request, 'faculty.html', {"msg": "⚠️ Please enter a valid email address!"})

    dbe, cmd = pool.ConnectionPool()
    
    try:
        cmd.execute(
            "SELECT COUNT(*) FROM Professor WHERE Professor_id = ?",
            [Professor_id]
        )
        count = cmd.fetchone()[0]
        
        if count > 0:
            dbe.close()
            return render(request, 'faculty.html', {"msg": f"⚠️ Faculty ID '{Professor_id}' already exists! Please use a different ID."})
        
        cmd.execute(
            "SELECT COUNT(*) FROM Courses WHERE course_code = ?",
            [course_code]
        )
        course_exists = cmd.fetchone()[0]
        
        if course_exists == 0:
            dbe.close()
            return render(request, 'faculty.html', {"msg": f"⚠️ Course '{course_code}' does not exist! Please add the course first."})
        
        cmd.execute(
            "INSERT INTO Professor (Professor_id, Professor_name, course_code, Phone_number, Email_ID) VALUES (?, ?, ?, ?, ?)",
            [Professor_id, Professor_name, course_code, str(s), Email_ID]
        )
        dbe.commit()
        msg = f"✅ Faculty '{Professor_name}' added successfully!"
            
    except Exception as e:
        dbe.rollback()
        error_msg = str(e)
        if "UNIQUE constraint failed" in error_msg:
            msg = f"⚠️ Faculty ID '{Professor_id}' already exists! Please use a different ID."
        elif "FOREIGN KEY constraint failed" in error_msg:
            msg = f"⚠️ Course '{course_code}' does not exist! Please add the course first."
        else:
            msg = f"❌ Error adding faculty: {error_msg}"
    finally:
        dbe.close()

    return render(request, 'faculty.html', {"msg": msg})

def generatepage(request):
    """ROBUST Timetable Generation Algorithm with MERGED Individual Faculty Timetables"""
    if request.method != 'POST':
        return HttpResponse(status=405)

    # Get input parameters
    program = request.POST.get('program', '').strip()
    semester = request.POST.get('semester', '').strip()
    batches = int(request.POST.get('batches') or 1)
    lecture_halls = int(request.POST.get('lecture_halls') or 1)
    absent_prof = request.POST.get('absent_prof', '').strip()
    replacement_prof = request.POST.get('replacement_prof', '').strip()

    # Store parameters in session
    request.session['last_program'] = program
    request.session['last_semester'] = semester
    request.session['last_batches'] = batches
    request.session['last_lecture_halls'] = lecture_halls
    request.session['last_absent'] = absent_prof
    request.session['last_replacement'] = replacement_prof

    dbe, cmd = pool.ConnectionPool()

    # Get programs list
    if program.upper() == 'ALL' or not program:
        cmd.execute("SELECT DISTINCT Program_name FROM Program ORDER BY Program_name")
        progs = [row[0] for row in cmd.fetchall()]
    else:
        progs = [program]

    # ============ GET ALL PROFESSOR DATA ============
    cmd.execute("""
        SELECT p.Professor_id, p.Professor_name, p.course_code 
        FROM Professor p
        ORDER BY p.Professor_id
    """)
    professors_raw = cmd.fetchall()
    
    # Build professor mappings
    professors_by_course = defaultdict(list)
    all_professors = {}
    professor_courses = defaultdict(list)
    
    for prof_id, prof_name, course_code in professors_raw:
        all_professors[prof_id] = prof_name
        if course_code:
            professors_by_course[course_code].append({
                'id': prof_id,
                'name': prof_name
            })
            professor_courses[prof_id].append(course_code)

    dbe.close()

    # ============ TIME SLOTS WITH LUNCH BREAK ============
    times = [
        "09:30-10:30",
        "10:30-11:30",
        "11:30-12:25",
        "12:25-01:15",
        "01:15-02:00",
        "02:00-03:30",
        "03:30-04:30",
        "04:30-05:30"
    ]
    
    days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday']
    
    # Calculate groups
    group_count = max(1, (batches + lecture_halls - 1) // lecture_halls) if lecture_halls > 0 else 1
    
    schedules = {}
    
    # NEW: Individual faculty timetable structure - MERGED BY NAME
    # Key will be professor NAME instead of ID to merge multiple IDs
    faculty_individual_timetables = defaultdict(lambda: {
        'name': '',
        'ids': [],  # Store all professor IDs for this name
        'schedule': [[{
            'course': 'Free',
            'course_name': '',
            'time_slot': times[idx],
            'day': day_name,
            'day_index': day_idx
        } for day_idx, day_name in enumerate(days)] for idx in range(len(times))],
        'courses_taught': [],  # List of unique courses
        'total_hours': 0
    })
    
    # ============ GENERATE TIMETABLE FOR EACH PROGRAM ============
    for prog in progs:
        # Get courses for this program and semester
        dbe, cmd = pool.ConnectionPool()
        cmd.execute("""
            SELECT c.course_code, c.course_name, c.frequency
            FROM Courses c
            WHERE LOWER(c.Program_name) = LOWER(?) AND c.semester_id = ?
            ORDER BY c.frequency DESC
        """, [prog, semester])
        courses = cmd.fetchall()
        dbe.close()

        if not courses:
            # No courses for this program/semester
            schedules[prog] = create_empty_timetable(times)
            continue

        # ============ PREPARE COURSE SLOTS ============
        lab_slots = []
        lecture_slots = []
        
        for course_code, course_name, frequency in courses:
            course_info = {
                'course': course_code,
                'name': course_name,
                'professors': professors_by_course.get(course_code, []),
                'is_lab': 'LAB' in course_code.upper() or 'LAB' in course_name.upper()
            }
            
            for _ in range(int(frequency or 0)):
                if course_info['is_lab']:
                    lab_slots.append(course_info.copy())
                else:
                    lecture_slots.append(course_info.copy())

        random.shuffle(lecture_slots)
        random.shuffle(lab_slots)

        # ============ GENERATE TIMETABLE FOR EACH GROUP ============
        for group in range(group_count):
            # Create timetable matrix
            timetable = [[{
                'course': 'Free',
                'prof': '',
                'prof_name': '',
                'is_lab': False
            } for _ in range(5)] for _ in range(8)]
            
            # Mark lunch period
            for day in range(5):
                timetable[4][day] = {
                    'course': 'LUNCH',
                    'prof': '',
                    'prof_name': '',
                    'is_lab': False
                }
            
            # Track professor workload
            professor_workload = defaultdict(int)
            professor_daily_load = defaultdict(lambda: defaultdict(int))
            professor_schedule = defaultdict(set)
            
            # Place labs
            lab_placements = []
            for course in lab_slots:
                placed = False
                attempts = 0
                max_attempts = 100
                
                while not placed and attempts < max_attempts:
                    time_slot = random.randint(5, 7)
                    day = random.randint(0, 4)
                    
                    if time_slot <= 6:
                        if (timetable[time_slot][day]['course'] == 'Free' and 
                            timetable[time_slot + 1][day]['course'] == 'Free'):
                            
                            available_prof = None
                            if course['professors']:
                                for prof in course['professors']:
                                    prof_id = prof['id']
                                    if (prof_id not in professor_schedule[(time_slot, day)] and
                                        prof_id not in professor_schedule[(time_slot + 1, day)]):
                                        if professor_daily_load[prof_id][day] <= 4:
                                            available_prof = prof
                                            break
                            
                            if available_prof or not course['professors']:
                                for offset in range(2):
                                    prof_id = available_prof['id'] if available_prof else ''
                                    prof_name = available_prof['name'] if available_prof else ''
                                    
                                    timetable[time_slot + offset][day] = {
                                        'course': course['course'],
                                        'prof': prof_id,
                                        'prof_name': prof_name,
                                        'is_lab': True
                                    }
                                    if prof_id and prof_name:
                                        professor_schedule[(time_slot + offset, day)].add(prof_id)
                                        professor_daily_load[prof_id][day] += 1
                                        professor_workload[prof_id] += 1
                                        
                                        # Update individual faculty timetable by NAME (merging)
                                        faculty_individual_timetables[prof_name]['name'] = prof_name
                                        if prof_id not in faculty_individual_timetables[prof_name]['ids']:
                                            faculty_individual_timetables[prof_name]['ids'].append(prof_id)
                                        
                                        # Only set if cell is free or this is a better assignment
                                        current_cell = faculty_individual_timetables[prof_name]['schedule'][time_slot + offset][day]
                                        if current_cell['course'] == 'Free' or current_cell['course'] == '':
                                            faculty_individual_timetables[prof_name]['schedule'][time_slot + offset][day] = {
                                                'course': course['course'],
                                                'course_name': course['name'],
                                                'prof': prof_id,
                                                'prof_name': prof_name,
                                                'is_lab': True,
                                                'time_slot': times[time_slot + offset],
                                                'day': days[day],
                                                'day_index': day
                                            }
                                            if course['course'] not in [c['code'] for c in faculty_individual_timetables[prof_name]['courses_taught']]:
                                                faculty_individual_timetables[prof_name]['courses_taught'].append({
                                                    'code': course['course'],
                                                    'name': course['name'],
                                                    'type': 'Lab'
                                                })
                                            faculty_individual_timetables[prof_name]['total_hours'] += 1
                                
                                placed = True
                    
                    attempts += 1
                
                if not placed:
                    lab_placements.append(course)

            # Place lectures
            course_index = 0
            lecture_list = lecture_slots.copy()
            random.shuffle(lecture_list)
            
            for time_slot in range(8):
                if time_slot == 4:
                    continue
                    
                for day in range(5):
                    if timetable[time_slot][day]['course'] != 'Free':
                        continue
                    
                    if course_index < len(lecture_list):
                        course = lecture_list[course_index]
                        
                        selected_prof = None
                        
                        if course['professors']:
                            best_score = -1
                            for prof in course['professors']:
                                prof_id = prof['id']
                                
                                if prof_id in professor_schedule[(time_slot, day)]:
                                    continue
                                
                                if professor_daily_load[prof_id][day] >= 5:
                                    continue
                                
                                score = professor_workload[prof_id]
                                
                                if best_score == -1 or score < best_score:
                                    best_score = score
                                    selected_prof = prof
                            
                            if not selected_prof and course['professors']:
                                min_load = float('inf')
                                for prof in course['professors']:
                                    load = professor_daily_load[prof['id']][day]
                                    if load < min_load:
                                        min_load = load
                                        selected_prof = prof
                        
                        if selected_prof or not course['professors']:
                            prof_id = selected_prof['id'] if selected_prof else ''
                            prof_name = selected_prof['name'] if selected_prof else ''
                            
                            timetable[time_slot][day] = {
                                'course': course['course'],
                                'prof': prof_id,
                                'prof_name': prof_name,
                                'is_lab': False
                            }
                            
                            if prof_id and prof_name:
                                professor_schedule[(time_slot, day)].add(prof_id)
                                professor_daily_load[prof_id][day] += 1
                                professor_workload[prof_id] += 1
                                
                                # Update individual faculty timetable by NAME (merging)
                                faculty_individual_timetables[prof_name]['name'] = prof_name
                                if prof_id not in faculty_individual_timetables[prof_name]['ids']:
                                    faculty_individual_timetables[prof_name]['ids'].append(prof_id)
                                
                                current_cell = faculty_individual_timetables[prof_name]['schedule'][time_slot][day]
                                if current_cell['course'] == 'Free' or current_cell['course'] == '':
                                    faculty_individual_timetables[prof_name]['schedule'][time_slot][day] = {
                                        'course': course['course'],
                                        'course_name': course['name'],
                                        'prof': prof_id,
                                        'prof_name': prof_name,
                                        'is_lab': False,
                                        'time_slot': times[time_slot],
                                        'day': days[day],
                                        'day_index': day
                                    }
                                    if course['course'] not in [c['code'] for c in faculty_individual_timetables[prof_name]['courses_taught']]:
                                        faculty_individual_timetables[prof_name]['courses_taught'].append({
                                            'code': course['course'],
                                            'name': course['name'],
                                            'type': 'Lecture'
                                        })
                                    faculty_individual_timetables[prof_name]['total_hours'] += 1
                            
                            course_index += 1

            # Handle labs that couldn't be placed
            for course in lab_placements:
                for time_slot in range(8):
                    if time_slot == 4:
                        continue
                    for day in range(5):
                        if timetable[time_slot][day]['course'] == 'Free':
                            timetable[time_slot][day] = {
                                'course': course['course'],
                                'prof': '',
                                'prof_name': 'TBD (Lab)',
                                'is_lab': True
                            }
                            break
                    else:
                        continue
                    break

            # Handle absent professor
            if absent_prof:
                timetable = handle_absent_professor(
                    timetable, absent_prof, replacement_prof, 
                    professors_by_course, all_professors, professor_schedule,
                    faculty_individual_timetables, times, days
                )

            # Optimize distribution
            timetable = optimize_timetable_distribution(timetable)

            # ============ FORMAT FOR DISPLAY - STRUCTURED DATA ============
            matrix_with_times = []
            for idx, row in enumerate(timetable):
                display_row = {
                    'time': times[idx],
                    'cells': []
                }
                for cell in row:
                    if cell['course'] == 'LUNCH':
                        display_row['cells'].append({
                            'type': 'lunch',
                            'display': '🍽️ LUNCH BREAK'
                        })
                    elif cell['course'] == 'Free':
                        display_row['cells'].append({
                            'type': 'free',
                            'display': 'Free'
                        })
                    elif cell['prof_name'] and cell['prof_name'] not in ['TBD', 'TBD (Lab)']:
                        course_name = get_course_name(cell['course'])
                        display_row['cells'].append({
                            'type': 'course',
                            'course': cell['course'],
                            'course_name': course_name,
                            'professor': cell['prof_name'],
                            'display': f"{cell['course']}\n{cell['prof_name']}"
                        })
                    elif cell['prof_name'] == 'TBD (Lab)':
                        course_name = get_course_name(cell['course'])
                        display_row['cells'].append({
                            'type': 'lab_needed',
                            'course': cell['course'],
                            'course_name': course_name,
                            'display': f"{cell['course']}\n(Lab Faculty Needed)"
                        })
                    elif cell['prof_name'] == 'TBD':
                        course_name = get_course_name(cell['course'])
                        display_row['cells'].append({
                            'type': 'faculty_needed',
                            'course': cell['course'],
                            'course_name': course_name,
                            'display': f"{cell['course']}\n(Faculty Needed)"
                        })
                    else:
                        course_name = get_course_name(cell['course'])
                        display_row['cells'].append({
                            'type': 'course',
                            'course': cell['course'],
                            'course_name': course_name,
                            'display': cell['course']
                        })
                matrix_with_times.append(display_row)
            
            group_name = prog if group_count == 1 else f"{prog} (Group {group+1})"
            schedules[group_name] = {'matrix': matrix_with_times}
    
    # Convert faculty_individual_timetables to regular dict for session
    faculty_timetables_dict = {}
    for prof_name, data in faculty_individual_timetables.items():
        if data['name']:
            # Convert schedule to serializable format
            serializable_schedule = []
            for time_idx, time_slot in enumerate(data['schedule']):
                serializable_row = []
                for day_idx, cell in enumerate(time_slot):
                    serializable_row.append({
                        'course': cell.get('course', 'Free'),
                        'course_name': cell.get('course_name', ''),
                        'prof': cell.get('prof', ''),
                        'prof_name': cell.get('prof_name', ''),
                        'is_lab': cell.get('is_lab', False),
                        'time_slot': times[time_idx] if time_idx < len(times) else '',
                        'day': days[day_idx] if day_idx < len(days) else '',
                        'day_index': day_idx
                    })
                serializable_schedule.append(serializable_row)
            
            # Use first ID as primary (or join multiple)
            primary_id = data['ids'][0] if data['ids'] else prof_name.replace(' ', '_')
            
            faculty_timetables_dict[primary_id] = {
                'name': data['name'],
                'id': primary_id,
                'all_ids': data['ids'],
                'schedule': serializable_schedule,
                'courses_taught': data['courses_taught'],
                'total_hours': data['total_hours']
            }
    
    request.session['schedule_results'] = schedules
    request.session['faculty_individual_timetables'] = faculty_timetables_dict
    
    return redirect('/output')

def handle_absent_professor(timetable, absent_prof, replacement_prof, professors_by_course, all_professors, professor_schedule, faculty_timetables=None, times=None, days=None):
    """Intelligently handle absent professor"""
    
    for time_slot in range(8):
        if time_slot == 4:
            continue
        for day in range(5):
            cell = timetable[time_slot][day]
            if cell['prof'] == absent_prof:
                
                if replacement_prof and replacement_prof in all_professors:
                    cell['prof'] = replacement_prof
                    cell['prof_name'] = all_professors[replacement_prof]
                    professor_schedule[(time_slot, day)].add(replacement_prof)
                    
                    # Update faculty timetable if provided
                    if faculty_timetables and times and days:
                        # Find by name
                        replacement_name = all_professors[replacement_prof]
                        if replacement_name in faculty_timetables:
                            faculty_timetables[replacement_name]['schedule'][time_slot][day] = {
                                'course': cell['course'],
                                'course_name': get_course_name(cell['course']),
                                'prof': replacement_prof,
                                'prof_name': replacement_name,
                                'is_lab': cell['is_lab'],
                                'time_slot': times[time_slot],
                                'day': days[day],
                                'day_index': day
                            }
                
                else:
                    course_code = cell['course']
                    if course_code in professors_by_course:
                        available_prof = None
                        for prof in professors_by_course[course_code]:
                            if prof['id'] != absent_prof:
                                if prof['id'] not in professor_schedule[(time_slot, day)]:
                                    available_prof = prof
                                    break
                        
                        if available_prof:
                            cell['prof'] = available_prof['id']
                            cell['prof_name'] = available_prof['name']
                            professor_schedule[(time_slot, day)].add(available_prof['id'])
                            
                            if faculty_timetables and times and days:
                                if available_prof['name'] in faculty_timetables:
                                    faculty_timetables[available_prof['name']]['schedule'][time_slot][day] = {
                                        'course': cell['course'],
                                        'course_name': get_course_name(cell['course']),
                                        'prof': available_prof['id'],
                                        'prof_name': available_prof['name'],
                                        'is_lab': cell['is_lab'],
                                        'time_slot': times[time_slot],
                                        'day': days[day],
                                        'day_index': day
                                    }
                        else:
                            cell['prof'] = ''
                            cell['prof_name'] = 'TBD'
                    else:
                        cell['prof'] = ''
                        cell['prof_name'] = 'TBD'
    
    return timetable

def optimize_timetable_distribution(timetable):
    """Ensure even distribution of courses across days"""
    
    day_counts = [0] * 5
    for time_slot in range(8):
        if time_slot == 4:
            continue
        for day in range(5):
            if timetable[time_slot][day]['course'] != 'Free' and timetable[time_slot][day]['course'] != 'LUNCH':
                day_counts[day] += 1
    
    if sum(day_counts) > 0:
        avg_courses_per_day = sum(day_counts) / 5
        
        for day in range(5):
            if day_counts[day] < avg_courses_per_day - 2:
                for other_day in range(5):
                    if day_counts[other_day] > avg_courses_per_day + 2:
                        for time_slot in range(8):
                            if time_slot == 4:
                                continue
                            if (timetable[time_slot][other_day]['course'] != 'Free' and
                                timetable[time_slot][other_day]['course'] != 'LUNCH' and
                                timetable[time_slot][day]['course'] == 'Free'):
                                
                                course_to_move = timetable[time_slot][other_day]
                                if not course_to_move['is_lab']:
                                    timetable[time_slot][day] = course_to_move
                                    timetable[time_slot][other_day] = {
                                        'course': 'Free',
                                        'prof': '',
                                        'prof_name': '',
                                        'is_lab': False
                                    }
                                    day_counts[day] += 1
                                    day_counts[other_day] -= 1
                                    break
                        break
    
    return timetable

def create_empty_timetable(times):
    """Create an empty timetable structure"""
    matrix = []
    for time in times:
        row = [time]
        for _ in range(5):
            row.append('Free')
        matrix.append(row)
    return {'matrix': matrix}

def replace_professor(request):
    """Handle professor replacement after timetable generation"""
    if request.method != 'POST':
        return HttpResponse(status=405)
    
    absent = request.POST.get('absent_prof', '').strip()
    replacement = request.POST.get('replacement_prof', '').strip()
    schedules = request.session.get('schedule_results', {})
    faculty_timetables = request.session.get('faculty_individual_timetables', {})
    
    if not schedules or not absent:
        return redirect('/output')
    
    dbe, cmd = pool.ConnectionPool()
    try:
        cmd.execute("SELECT Professor_id, Professor_name, course_code FROM Professor")
        professors = cmd.fetchall()
    except Exception as e:
        professors = []
        print(f"Error fetching professors: {e}")
    finally:
        dbe.close()
    
    prof_dict = {p[0]: p[1] for p in professors}
    
    # Update main timetables
    for prog, data in schedules.items():
        matrix = data['matrix']
        for i in range(len(matrix)):
            row = matrix[i]
            for j in range(1, len(row)):
                cell = row[j]
                if isinstance(cell, dict) and cell.get('type') == 'course':
                    if cell.get('professor') == absent:
                        if replacement and replacement in prof_dict:
                            cell['professor'] = prof_dict[replacement]
                            cell['display'] = f"{cell['course']}\n{prof_dict[replacement]}"
                        else:
                            cell['type'] = 'faculty_needed'
                            cell['display'] = f"{cell['course']}\n(Faculty Needed)"
    
    request.session['schedule_results'] = schedules
    request.session['faculty_individual_timetables'] = faculty_timetables
    return redirect('/output')

def download_faculty_individual_timetable(request, faculty_id):
    """Download individual faculty timetable in requested format"""
    if request.method != 'POST':
        return HttpResponse(status=405)
    
    fmt = request.POST.get('format', 'pdf').lower()
    faculty_timetables = request.session.get('faculty_individual_timetables', {})
    
    if faculty_id not in faculty_timetables:
        return HttpResponse('Faculty timetable not found', status=400)
    
    faculty = faculty_timetables[faculty_id]
    
    # Create HTML for download
    times = ["09:30-10:30", "10:30-11:30", "11:30-12:25", "12:25-01:15", "01:15-02:00", "02:00-03:30", "03:30-04:30", "04:30-05:30"]
    days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday']
    
    html = f"""<!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>{faculty['name']} - Individual Timetable</title>
        <style>
            body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 40px; background: white; color: #333; }}
            .header {{ text-align: center; margin-bottom: 40px; border-bottom: 2px solid #ff0000; padding-bottom: 20px; }}
            .college-name {{ font-size: 24px; font-weight: bold; color: #ff0000; margin-bottom: 10px; }}
            .title {{ font-size: 20px; font-weight: 600; margin: 10px 0; }}
            .faculty-name {{ font-size: 18px; margin: 10px 0; }}
            table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
            th {{ background: #ff0000; color: white; padding: 12px; border: 1px solid #ddd; }}
            td {{ padding: 10px; border: 1px solid #ddd; text-align: center; }}
            .time-col {{ background: #f5f5f5; font-weight: bold; }}
            .lunch-cell {{ background: #fff3cd; color: #856404; }}
            .free-cell {{ color: #999; font-style: italic; }}
            .lab-cell {{ background: #e8f4f8; }}
            .footer {{ margin-top: 40px; text-align: center; font-size: 12px; color: #666; border-top: 1px solid #eee; padding-top: 20px; }}
        </style>
    </head>
    <body>
        <div class="header">
            <div class="college-name">DELHI TECHNICAL CAMPUS, Greater Noida</div>
            <div class="title">Time Table (Even Session 2026-27)</div>
            <div class="faculty-name">{faculty['name']} Individual Time Table</div>
        </div>
        
        <table>
            <thead>
                <tr>
                    <th>Time</th>
                    <th>Monday</th>
                    <th>Tuesday</th>
                    <th>Wednesday</th>
                    <th>Thursday</th>
                    <th>Friday</th>
                </thead>
            <tbody>"""
    
    for time_idx, time_slot in enumerate(times):
        html += f"""<tr>
            <td class="time-col">{time_slot}</td>"""
        for day_idx in range(5):
            if time_idx < len(faculty['schedule']) and day_idx < len(faculty['schedule'][time_idx]):
                cell = faculty['schedule'][time_idx][day_idx]
            else:
                cell = {'course': 'Free', 'course_name': '', 'is_lab': False}
            
            course = cell.get('course', 'Free')
            course_name = cell.get('course_name', '')
            is_lab = cell.get('is_lab', False)
            
            if course == 'LUNCH':
                html += f'<td class="lunch-cell">🍽️ LUNCH BREAK</td>'
            elif course == 'Free' or course == '':
                html += f'<td class="free-cell">—</td>'
            else:
                lab_class = 'lab-cell' if is_lab else ''
                html += f'''<td class="{lab_class}">
                    <strong>{course}</strong><br>
                    <small>{course_name}</small>
                </td>'''
        html += '</tr>'
    
    html += f"""
            </tbody>
        </table>
        
        <div class="footer">
            <p>Generated by TEMPOSYNK | Total Teaching Hours: {faculty['total_hours']} hrs/week</p>
            <p>Courses: {', '.join([c['code'] + ' (' + c['type'] + ')' for c in faculty['courses_taught']])}</p>
        </div>
    </body>
    </html>"""
    
    # Generate file based on format
    if fmt in ('csv', 'text/csv'):
        return generate_csv_from_faculty_timetable(faculty, times, days)
    elif fmt in ('xls', 'xlsx', 'excel'):
        return generate_excel_from_faculty_timetable(faculty, times, days)
    elif fmt in ('doc', 'docx', 'word'):
        return generate_word_from_faculty_timetable(faculty, times, days)
    else:
        # PDF
        pdf = io.BytesIO()
        try:
            pisa_status = pisa.CreatePDF(src=html, dest=pdf)
        except Exception:
            return HttpResponse('PDF generation failed', status=500)
        
        if pisa_status.err:
            return HttpResponse('Error generating PDF', status=500)
        
        pdf.seek(0)
        filename = f"{faculty['name'].replace(' ', '_')}_Timetable.pdf"
        return FileResponse(pdf, as_attachment=True, filename=filename, content_type='application/pdf')

def generate_csv_from_faculty_timetable(faculty, times, days):
    """Generate CSV from faculty timetable"""
    import csv
    output = io.StringIO()
    writer = csv.writer(output)
    
    writer.writerow([f"DELHI TECHNICAL CAMPUS, Greater Noida"])
    writer.writerow([f"Time Table (Even Session 2026-27)"])
    writer.writerow([f"{faculty['name']} Individual Time Table"])
    writer.writerow([])
    
    header = ['Time'] + days
    writer.writerow(header)
    
    for time_idx, time_slot in enumerate(times):
        row = [time_slot]
        for day_idx in range(5):
            if time_idx < len(faculty['schedule']) and day_idx < len(faculty['schedule'][time_idx]):
                cell = faculty['schedule'][time_idx][day_idx]
            else:
                cell = {'course': 'Free'}
            course = cell.get('course', 'Free')
            if course == 'LUNCH':
                row.append('LUNCH BREAK')
            elif course == 'Free' or course == '':
                row.append('Free')
            else:
                row.append(course)
        writer.writerow(row)
    
    output.seek(0)
    response = HttpResponse(output.getvalue(), content_type='text/csv')
    filename = f"{faculty['name'].replace(' ', '_')}_Timetable.csv"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response

def generate_excel_from_faculty_timetable(faculty, times, days):
    """Generate Excel from faculty timetable"""
    import pandas as pd
    data = []
    for time_idx, time_slot in enumerate(times):
        row = {'Time': time_slot}
        for day_idx, day in enumerate(days):
            if time_idx < len(faculty['schedule']) and day_idx < len(faculty['schedule'][time_idx]):
                cell = faculty['schedule'][time_idx][day_idx]
            else:
                cell = {'course': 'Free'}
            course = cell.get('course', 'Free')
            if course == 'LUNCH':
                row[day] = 'LUNCH BREAK'
            elif course == 'Free' or course == '':
                row[day] = 'Free'
            else:
                row[day] = course
        data.append(row)
    
    df = pd.DataFrame(data)
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name=faculty['name'][:31])
    bio.seek(0)
    filename = f"{faculty['name'].replace(' ', '_')}_Timetable.xlsx"
    return FileResponse(bio, as_attachment=True, filename=filename, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

def generate_word_from_faculty_timetable(faculty, times, days):
    """Generate Word document from faculty timetable"""
    doc = Document()
    doc.add_heading('DELHI TECHNICAL CAMPUS, Greater Noida', 0)
    doc.add_heading('Time Table (Even Session 2026-27)', 1)
    doc.add_heading(f"{faculty['name']} Individual Time Table", 2)
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Time'
    for i, day in enumerate(days):
        hdr_cells[i + 1].text = day
    
    for time_idx, time_slot in enumerate(times):
        row_cells = table.add_row().cells
        row_cells[0].text = time_slot
        for day_idx in range(5):
            if time_idx < len(faculty['schedule']) and day_idx < len(faculty['schedule'][time_idx]):
                cell = faculty['schedule'][time_idx][day_idx]
            else:
                cell = {'course': 'Free'}
            course = cell.get('course', 'Free')
            if course == 'LUNCH':
                row_cells[day_idx + 1].text = 'LUNCH BREAK'
            elif course == 'Free' or course == '':
                row_cells[day_idx + 1].text = 'Free'
            else:
                row_cells[day_idx + 1].text = course
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    filename = f"{faculty['name'].replace(' ', '_')}_Timetable.docx"
    return FileResponse(bio, as_attachment=True, filename=filename, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

def download_timetable(request):
    """Download main timetable in various formats"""
    if request.method != 'POST':
        return HttpResponse(status=405)
    
    fmt = request.POST.get('format', 'pdf').lower()
    html = request.POST.get('html', '')
    person_name = request.POST.get('person_name', 'User')
    college_name = request.POST.get('college_name', 'College')
    logo_data = request.POST.get('logo_data', '')
    
    if not html:
        return HttpResponse('No HTML provided', status=400)
    
    try:
        dfs = pd.read_html(html)
    except Exception:
        dfs = []
    
    if not dfs:
        return HttpResponse('No table found', status=400)
    
    df = dfs[0]
    
    if fmt in ('csv', 'text/csv'):
        bio = io.BytesIO()
        bio.write(df.to_csv(index=False).encode('utf-8'))
        bio.seek(0)
        return FileResponse(bio, as_attachment=True, filename='timetable.csv', content_type='text/csv')
    
    if fmt in ('xls', 'xlsx', 'excel'):
        bio = io.BytesIO()
        with pd.ExcelWriter(bio, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Timetable')
        bio.seek(0)
        return FileResponse(bio, as_attachment=True, filename='timetable.xlsx', 
                          content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    
    if fmt in ('doc', 'docx', 'word'):
        doc = Document()
        doc.add_heading(f'{college_name}', 0)
        doc.add_heading('TEMPOSYNK - Generated Timetable', 1)
        doc.add_paragraph(f'Generated for: {person_name}')
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph('')
        
        table = doc.add_table(rows=1, cols=len(df.columns))
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)
        
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
        
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return FileResponse(bio, as_attachment=True, filename='timetable.docx',
                          content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
    pdf = io.BytesIO()
    try:
        pisa_status = pisa.CreatePDF(src=html, dest=pdf)
    except Exception:
        return HttpResponse('PDF generation failed', status=500)
    
    if pisa_status.err:
        return HttpResponse('Error generating PDF', status=500)
    
    pdf.seek(0)
    return FileResponse(pdf, as_attachment=True, filename='timetable.pdf', content_type='application/pdf')

# View functions
def callview_program(request):
    dbe, cmd = pool.ConnectionPool()
    try:
        cmd.execute("SELECT Program_name, semester_id FROM Program ORDER BY Program_name, CAST(semester_id AS INTEGER)")
        programs = cmd.fetchall()
    except Exception as e:
        programs = []
        print(f"Error fetching programs: {e}")
    finally:
        dbe.close()
    return render(request, 'view_program.html', {"programs": programs})

def callview_course(request):
    dbe, cmd = pool.ConnectionPool()
    try:
        cmd.execute("SELECT * FROM Courses ORDER BY Program_name, CAST(semester_id AS INTEGER), course_code")
        courses = cmd.fetchall()
    except Exception as e:
        courses = []
        print(f"Error fetching courses: {e}")
    finally:
        dbe.close()
    return render(request, 'view_course.html', {"courses": courses})

def callview_lectureRoom(request):
    dbe, cmd = pool.ConnectionPool()
    try:
        cmd.execute("SELECT * FROM Lecture ORDER BY lecture_id")
        lectures = cmd.fetchall()
    except Exception as e:
        lectures = []
        print(f"Error fetching lecture rooms: {e}")
    finally:
        dbe.close()
    return render(request, 'view_lectureRoom.html', {"lectures": lectures})

def callview_labRoom(request):
    dbe, cmd = pool.ConnectionPool()
    try:
        cmd.execute("SELECT * FROM Lab ORDER BY Lab_id")
        labs = cmd.fetchall()
    except Exception as e:
        labs = []
        print(f"Error fetching labs: {e}")
    finally:
        dbe.close()
    return render(request, 'view_labRoom.html', {"labs": labs})

def callview_faculty(request):
    dbe, cmd = pool.ConnectionPool()
    try:
        cmd.execute("SELECT * FROM Professor ORDER BY Professor_id")
        professors = cmd.fetchall()
    except Exception as e:
        professors = []
        print(f"Error fetching faculty: {e}")
    finally:
        dbe.close()
    return render(request, 'view_faculty.html', {"Professors": professors})

# Delete functions
def delete_course(request):
    cid = request.GET.get('cid', '').strip()
    if not cid:
        return redirect('/callview_course')
    
    dbe, cmd = pool.ConnectionPool()
    try:
        cmd.execute("SELECT COUNT(*) FROM Professor WHERE course_code = ?", [cid])
        prof_count = cmd.fetchone()[0]
        
        if prof_count > 0:
            cmd.execute("DELETE FROM Professor WHERE course_code = ?", [cid])
        
        cmd.execute("DELETE FROM Courses WHERE course_code = ?", [cid])
        dbe.commit()
        msg = f"✅ Course '{cid}' deleted successfully!"
    except Exception as e:
        dbe.rollback()
        msg = f"❌ Error deleting course: {str(e)}"
    finally:
        cmd.execute("SELECT * FROM Courses")
        courses = cmd.fetchall()
        dbe.close()
    
    return render(request, 'view_course.html', {'courses': courses, 'msg': msg})

def delete_program(request):
    """Delete a specific program for a specific semester only"""
    program_name = request.GET.get('program_name', '').strip()
    semester_id = request.GET.get('semester_id', '').strip()
    prog_id = request.GET.get('id', '').strip()
    
    dbe, cmd = pool.ConnectionPool()
    msg = ""
    
    try:
        if program_name and semester_id:
            cmd.execute(
                "SELECT COUNT(*) FROM Program WHERE Program_name = ? AND semester_id = ?",
                [program_name, semester_id]
            )
            exists = cmd.fetchone()[0]
            
            if exists == 0:
                msg = f"⚠️ Program '{program_name}' for Semester {semester_id} does not exist!"
            else:
                cmd.execute(
                    "DELETE FROM Courses WHERE Program_name = ? AND semester_id = ?",
                    [program_name, semester_id]
                )
                try:
                    deleted_courses = cmd.rowcount
                except AttributeError:
                    cmd.execute(
                        "SELECT COUNT(*) FROM Courses WHERE Program_name = ? AND semester_id = ?",
                        [program_name, semester_id]
                    )
                    deleted_courses = cmd.fetchone()[0]
                
                cmd.execute(
                    "DELETE FROM Program WHERE Program_name = ? AND semester_id = ?",
                    [program_name, semester_id]
                )
                dbe.commit()
                msg = f"✅ Program '{program_name}' for Semester {semester_id} deleted successfully! ({deleted_courses} course(s) removed)"
                
        elif prog_id and '_' in prog_id:
            parts = prog_id.split('_')
            if len(parts) == 2:
                prog_name = parts[0]
                sem_id = parts[1]
                
                cmd.execute(
                    "SELECT COUNT(*) FROM Program WHERE Program_name = ? AND semester_id = ?",
                    [prog_name, sem_id]
                )
                exists = cmd.fetchone()[0]
                
                if exists == 0:
                    msg = f"⚠️ Program '{prog_name}' for Semester {sem_id} does not exist!"
                else:
                    cmd.execute(
                        "DELETE FROM Courses WHERE Program_name = ? AND semester_id = ?",
                        [prog_name, sem_id]
                    )
                    try:
                        deleted_courses = cmd.rowcount
                    except AttributeError:
                        cmd.execute(
                            "SELECT COUNT(*) FROM Courses WHERE Program_name = ? AND semester_id = ?",
                            [prog_name, sem_id]
                        )
                        deleted_courses = cmd.fetchone()[0]
                    
                    cmd.execute(
                        "DELETE FROM Program WHERE Program_name = ? AND semester_id = ?",
                        [prog_name, sem_id]
                    )
                    dbe.commit()
                    msg = f"✅ Program '{prog_name}' for Semester {sem_id} deleted successfully! ({deleted_courses} course(s) removed)"
            else:
                msg = "⚠️ Invalid program identifier format"
        else:
            msg = "⚠️ Please specify both program name and semester"
            
    except Exception as e:
        dbe.rollback()
        msg = f"❌ Error deleting program: {str(e)}"
        print(f"Delete program error: {e}")
    
    try:
        cmd.execute("SELECT Program_name, semester_id FROM Program ORDER BY Program_name, CAST(semester_id AS INTEGER)")
        programs = cmd.fetchall()
    except Exception as e:
        programs = []
        print(f"Error fetching programs: {e}")
    finally:
        dbe.close()
    
    return render(request, 'view_program.html', {"programs": programs, "msg": msg})

def delete_faculty(request):
    fid = request.GET.get('id', '').strip()
    if not fid:
        return redirect('/callview_faculty')
    
    dbe, cmd = pool.ConnectionPool()
    try:
        cmd.execute("DELETE FROM Professor WHERE Professor_id = ?", [fid])
        dbe.commit()
        msg = f"✅ Faculty '{fid}' deleted successfully!"
    except Exception as e:
        dbe.rollback()
        msg = f"❌ Error deleting faculty: {str(e)}"
    finally:
        cmd.execute("SELECT * FROM Professor")
        professors = cmd.fetchall()
        dbe.close()
    
    return render(request, 'view_faculty.html', {"Professors": professors, "msg": msg})

def delete_lectureroom(request):
    lid = request.GET.get('id', '').strip()
    if not lid:
        return redirect('/callview_lectureRoom')
    
    dbe, cmd = pool.ConnectionPool()
    try:
        cmd.execute("DELETE FROM Lecture WHERE lecture_id = ?", [lid])
        dbe.commit()
        msg = f"✅ Lecture room '{lid}' deleted successfully!"
    except Exception as e:
        dbe.rollback()
        msg = f"❌ Error deleting lecture room: {str(e)}"
    finally:
        cmd.execute("SELECT * FROM Lecture")
        lectures = cmd.fetchall()
        dbe.close()
    
    return render(request, 'view_lectureRoom.html', {"lectures": lectures, "msg": msg})

def delete_labroom(request):
    lid = request.GET.get('id', '').strip()
    if not lid:
        return redirect('/callview_labRoom')
    
    dbe, cmd = pool.ConnectionPool()
    try:
        cmd.execute("DELETE FROM Lab WHERE Lab_id = ?", [lid])
        dbe.commit()
        msg = f"✅ Lab room '{lid}' deleted successfully!"
    except Exception as e:
        dbe.rollback()
        msg = f"❌ Error deleting lab room: {str(e)}"
    finally:
        cmd.execute("SELECT * FROM Lab")
        labs = cmd.fetchall()
        dbe.close()
    
    return render(request, 'view_labRoom.html', {"labs": labs, "msg": msg})